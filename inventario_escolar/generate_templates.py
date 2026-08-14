import os
from docx import Document

# La carpeta de plantillas se resuelve relativa a este archivo (raíz del proyecto)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates_word')


def create_template(filename, title):
    doc = Document()
    doc.add_heading(f'Formato de {title}', 0)

    doc.add_paragraph('Datos del Alumno:')
    doc.add_paragraph('Nombre: {{ nombre }}')
    doc.add_paragraph('Matrícula: {{ matricula }}')
    doc.add_paragraph('Carrera: {{ carrera }}')
    doc.add_paragraph('Generación: {{ generacion }}')
    doc.add_paragraph('Estatus: {{ estatus }}')

    doc.add_heading('Información del Expediente', level=1)
    doc.add_paragraph('Expediente Base: {{ expediente_base }}')
    doc.add_paragraph('Clave del Expediente: {{ clave_expediente }}')
    doc.add_paragraph('Tipo de Módulo: {{ tipo_modulo }}')
    doc.add_paragraph('Sector: {{ sector }}')
    doc.add_paragraph('Dependencia / Lugar: {{ dependencia }}')

    doc.add_paragraph('\n\nFecha de emisión: {{ fecha }}')

    path = os.path.join(TEMPLATES_DIR, filename)
    doc.save(path)
    print(f'Created {path}')


if __name__ == '__main__':
    # Ensure directory exists
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    create_template('plantilla_practicas.docx', 'Prácticas Profesionales')
    create_template('plantilla_servicio.docx', 'Servicio Social')
    create_template('plantilla_vinculacion.docx', 'Vinculación')
    print(f'Plantillas generadas en: {TEMPLATES_DIR}')

